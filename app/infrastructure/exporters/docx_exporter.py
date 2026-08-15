"""
Применяет принятые правки к docx: абзац с вхождением old_text переписывается целиком
(все runs очищаются и заменяются одним run с новым текстом — внутреннее форматирование
изменённого абзаца не сохраняется, это упрощение MVP).
"""
import io

from docx import Document as DocxDocument

from app.domain.interfaces.document_exporter import AppliedChange


def _set_paragraph_text(paragraph, new_text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


class DocxExporter:
    def apply_changes(self, raw_bytes: bytes, changes: list[AppliedChange]) -> bytes:
        document = DocxDocument(io.BytesIO(raw_bytes))
        for change in changes:
            if change.change_type in ("modify", "delete") and change.old_text:
                for paragraph in document.paragraphs:
                    if change.old_text in paragraph.text:
                        new_text = paragraph.text.replace(change.old_text, change.new_text or "", 1)
                        _set_paragraph_text(paragraph, new_text)
                        break
            elif change.change_type == "add" and change.new_text:
                document.add_paragraph(change.new_text)
        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()
