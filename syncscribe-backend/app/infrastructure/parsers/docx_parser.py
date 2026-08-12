import io
from docx import Document as DocxDocument
from app.domain.interfaces.document_parser import DocumentSection, ParsedDocument


class DocxParser:
    def parse(self, raw_bytes: bytes) -> ParsedDocument:
        docx_document = DocxDocument(io.BytesIO(raw_bytes))
        sections: list[DocumentSection] = []
        text_parts: list[str] = []
        offset = 0
        for paragraph in docx_document.paragraphs:
            paragraph_text = paragraph.text
            if paragraph.style is not None and paragraph.style.name.startswith("Heading"):
                sections.append(DocumentSection(ref=paragraph_text.strip() or "Без названия", start_offset=offset, end_offset=offset))
            text_parts.append(paragraph_text)
            offset += len(paragraph_text) + 1
        plain_text = "\n".join(text_parts)
        for i, section in enumerate(sections):
            section.end_offset = sections[i + 1].start_offset if i + 1 < len(sections) else len(plain_text)
        return ParsedDocument(plain_text=plain_text, sections=sections)
